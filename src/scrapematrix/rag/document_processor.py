"""Document processing for RAG system."""
import logging
from pathlib import Path
from typing import List, Optional
import uuid

from .knowledge_base import Document, Chunk

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes documents into chunks for RAG system."""
    
    def __init__(self, chunk_size: int = 500, overlap: int = 100):
        """Initialize document processor.
        
        Args:
            chunk_size: Size of chunks in characters
            overlap: Number of overlapping characters between chunks
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def process_file(self, file_path: Path, title: Optional[str] = None) -> tuple[Document, List[Chunk]]:
        """Process a file into document and chunks.
        
        Args:
            file_path: Path to file
            title: Optional title for document
            
        Returns:
            Tuple of (Document, List[Chunk])
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Read file content
        content = self._read_file(file_path)
        
        # Create document
        doc_id = str(uuid.uuid4())
        doc_title = title or file_path.stem
        doc = Document(
            id=doc_id,
            title=doc_title,
            content=content,
            file_type=file_path.suffix.lower().lstrip('.'),
            file_path=str(file_path),
            metadata={
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
            }
        )
        
        # Create chunks
        chunks = self._chunk_text(doc_id, content)
        
        logger.info(f"Processed file {file_path.name}: {len(chunks)} chunks created")
        
        return doc, chunks
    
    def _read_file(self, file_path: Path) -> str:
        """Read content from file.
        
        Args:
            file_path: Path to file
            
        Returns:
            File content as string
        """
        file_type = file_path.suffix.lower()
        
        if file_type == '.txt':
            return file_path.read_text(encoding='utf-8', errors='ignore')
        
        elif file_type == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''.join(page.extract_text() for page in reader.pages)
                return text
            except ImportError:
                logger.warning("PyPDF2 not installed, returning empty content")
                return ""
        
        elif file_type in ['.docx', '.doc']:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = '\n'.join(para.text for para in doc.paragraphs)
                return text
            except ImportError:
                logger.warning("python-docx not installed, returning empty content")
                return ""
        
        elif file_type == '.md':
            return file_path.read_text(encoding='utf-8', errors='ignore')
        
        elif file_type == '.csv':
            return file_path.read_text(encoding='utf-8', errors='ignore')
        
        else:
            # Try reading as text
            try:
                return file_path.read_text(encoding='utf-8', errors='ignore')
            except Exception as e:
                logger.error(f"Could not read file {file_path}: {e}")
                return ""
    
    def _chunk_text(self, doc_id: str, text: str) -> List[Chunk]:
        """Split text into chunks with overlap.
        
        Args:
            doc_id: Document ID
            text: Text to chunk
            
        Returns:
            List of chunks
        """
        chunks = []
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_start = 0
        chunk_index = 0
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            # If adding this paragraph exceeds chunk size, save current chunk
            if len(current_chunk) + len(para) > self.chunk_size and current_chunk:
                chunk = Chunk(
                    id=f"{doc_id}_chunk_{chunk_index}",
                    document_id=doc_id,
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                )
                chunks.append(chunk)
                
                # Move back by overlap amount
                overlap_text = current_chunk[-self.overlap:] if len(current_chunk) > self.overlap else ""
                current_chunk = overlap_text + para
                current_start = max(0, current_start + len(current_chunk) - len(overlap_text) - len(para))
                chunk_index += 1
            else:
                current_chunk += "\n\n" + para if current_chunk else para
        
        # Add final chunk
        if current_chunk.strip():
            chunk = Chunk(
                id=f"{doc_id}_chunk_{chunk_index}",
                document_id=doc_id,
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
            )
            chunks.append(chunk)
        
        return chunks
    
    def process_text(self, text: str, title: str = "Untitled", file_type: str = "txt") -> tuple[Document, List[Chunk]]:
        """Process raw text into document and chunks.
        
        Args:
            text: Text content
            title: Document title
            file_type: File type identifier
            
        Returns:
            Tuple of (Document, List[Chunk])
        """
        doc_id = str(uuid.uuid4())
        
        doc = Document(
            id=doc_id,
            title=title,
            content=text,
            file_type=file_type,
        )
        
        chunks = self._chunk_text(doc_id, text)
        
        logger.info(f"Processed text: {len(chunks)} chunks created")
        
        return doc, chunks
